# -*- coding: utf-8 -*-
import os, zipfile, sys

# ==== הגדרות נתיבים ====
OUT_DIR = os.path.join(os.getcwd(), "kamilahome_package_output")
os.makedirs(OUT_DIR, exist_ok=True)

DOCX_PATH = os.path.join(OUT_DIR, "אפיון_קמילהום.docx")
PDF_PATH  = os.path.join(OUT_DIR, "אפיון_קמילהום.pdf")

# קבצי ויירפריימים
WIREFRAMES = [
    ("wireframe_dashboard.png", "Wireframe – Dashboard"),
    ("wireframe_event_form.png", "Wireframe – Event Form"),
    ("wireframe_customers_list.png", "Wireframe – Customers List"),
    ("wireframe_customer_form.png", "Wireframe – Customer Form"),
    ("wireframe_packages_catalog.png", "Wireframe – Packages Catalog"),
    ("wireframe_contract_templates.png", "Wireframe – Contract Templates"),
    ("wireframe_contract_preview.png", "Wireframe – Contract Preview & Sign"),
    ("wireframe_reports.png", "Wireframe – Reports"),
    ("wireframe_settings_roles.png", "Wireframe – Settings & Roles"),
    ("wireframe_events_list.png", "Wireframe – Events List"),
    ("wireframe_event_details.png", "Wireframe – Event Details"),
    ("wireframe_payments.png", "Wireframe – Payments"),
    ("wireframe_calendar.png", "Wireframe – Calendar View"),
]

ZIP_PATH = os.path.join(OUT_DIR, "אפיון_קמילהום_חבילה_מעודכן.zip")

# ==== תוכן המסמך ====
SPEC_SECTIONS = [
    ("1) תרשים זרימה מקצה לקצה",
     "A. ליד/לקוח → B. פתיחת אירוע → C. בניית הצעה/חבילה → D. חישוב עלויות → E. חוזה → F. חתימה → G. הפקת מסמכים/שליחה → H. סטטוס/גבייה → I. סגירת אירוע ודוחות\n\n"
     "1. קליטת לקוח/ליד\n   - הזנת פרטי קשר → בדיקת כפילויות (טלפון/אימייל) → שמירה.\n"
     "2. פתיחת אירוע\n   - שיוך ללקוח → תאריך/שעה/מיקום/סוג אירוע → סטטוס ראשוני = 'טיוטה'.\n"
     "3. בחירת חבילה\n   - בחירת חבילת בסיס + תוספות → מחיר מנה למבוגר/ילד → מדיניות הנחה.\n"
     "4. חישוב עלות\n   - חישוב אוטומטי + מע\"מ + הנחות/דמי מקדמה → הצגת סיכום.\n"
     "5. הפקת חוזה\n   - מיזוג תבנית → PDF → תצוגה מקדימה.\n"
     "6. חתימה דיגיטלית\n   - שליחה ללקוח לחתימה → קליטת מסמך חתום → עדכון סטטוס = 'חתום/ממתין לתשלום'.\n"
     "7. חשבונית/קבלה (אופציונלי לאינטגרציה עתידית)\n   - יצירת דרישת תשלום/קישור סליקה.\n"
     "8. ניהול סטטוסים\n   - טיוטה → בהצעה → ממתין לחתימה → חתום → מאושר/שולם חלקית/שולם מלא → בוצע → סגור/מבוטל.\n"
     "9. דוחות\n   - הכנסות לפי חודש/חבילה/סוג אירוע, המרות ליד→חוזה, תזרים צפוי."),
    ("2) סטטוסי אירוע ומעברים",
     "טבלת סטטוסים, פעולות מותרות, טריגרים למעברים, וחוקי ולידציה (חתימה מחייבת פרטי לקוח מלאים וכו')."),
    ("3) אפיון מסכים (UI)",
     "לוח מחוונים (Dashboard), ניהול לקוחות, ניהול אירועים, ניהול חבילות, חוזים והחתמה, דוחות, הרשאות והגדרות."),
    ("4) שדות, ולידציות, וכללי חישוב",
     "חישוב מחיר אירוע, כללי כפילויות, שדות חובה."),
    ("5) אינטגרציות (MVP → עתידי)",
     "מיילים/SMTP, חתימה דיגיטלית, סליקה/חשבוניות, Calendar iCal."),
    ("6) UX מיקרו־אינטרקציות",
     "Wizard, שמירה אוטומטית, Toast, מצב קריאה בלבד לאחר תשלום מלא."),
    ("7) API (סקיצה ל-Backend)",
     "Endpoints: customers, events, quote, contract, sign, payment, reports. אבטחה: JWT, RBAC, rate limits."),
    ("8) דוגמאות תבניות",
     "תבנית חוזה ותבנית מייל עם משתני Merge."),
    ("9) דוחות",
     "הכנסות חודשיות, המרות, גבייה פתוחה."),
    ("10) Non‑Functional",
     "ביצועים (pagination, אינדקסים), אבטחת מידע (הצפנה, audit), גיבוי, נגישות (WCAG)."),
    ("11) Backlog (MVP → V1)",
     "MVP: לקוחות, אירועים, חבילות, חוזה PDF, חתימה/שליחה, דוחות בסיס.\n"
     "V1: סליקה/חשבוניות, iCal, מובייל, התראות WhatsApp, פורטל לקוח."),
]

# ==== 1) יצירת DOCX ====
def build_docx():
    try:
        from docx import Document
    except ImportError:
        print("❌ חסרה חבילה: python-docx. התקן עם: pip install python-docx")
        sys.exit(1)

    doc = Document()
    doc.add_heading("קמילהום – תרשים זרימה ואפיון UI מפורט", level=1)
    doc.add_paragraph("גרסה: 0.1 • תאריך: 22/08/2025")
    doc.add_paragraph("---")

    for title, content in SPEC_SECTIONS:
        doc.add_heading(title, level=2)
        for para in content.split("\n"):
            doc.add_paragraph(para)

    doc.save(DOCX_PATH)
    print(f"✅ נוצר DOCX: {DOCX_PATH}")

# ==== 2) יצירת PDF (ReportLab) ====
def build_pdf():
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_RIGHT
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError:
        print("❌ חסרה חבילה: reportlab. התקן עם: pip install reportlab")
        sys.exit(1)

    # ניסיון לרשום גופן עם תמיכה טובה בעברית (אם מותקן במערכת)
    # מומלץ להוריד/להתקין DejaVu Sans או Rubik ולהציב כאן את הנתיב המלא אם צריך.
    preferred_fonts = [
        # הוסף נתיב מלא אם ברצונך להבטיח עברית תקינה:
        # r"C:\Windows\Fonts\Rubik-Regular.ttf",
        # "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    font_registered = False
    for fpath in preferred_fonts:
        if os.path.exists(fpath):
            try:
                pdfmetrics.registerFont(TTFont("HebrewUI", fpath))
                base_font = "HebrewUI"
                font_registered = True
                break
            except Exception:
                pass
    if not font_registered:
        # fallback: Helvetica (ייתכן שתצוגת עברית תהיה הפוכה/חסרה בחלק מהמערכות)
        base_font = "Helvetica"

    styles = getSampleStyleSheet()
    rtl_style = ParagraphStyle(
        'RTL',
        parent=styles['Normal'],
        fontName=base_font,
        fontSize=11,
        leading=15,
        alignment=TA_RIGHT,
    )
    rtl_heading = ParagraphStyle(
        'RTLHeading',
        parent=styles['Heading2'],
        fontName=base_font,
        fontSize=13,
        leading=16,
        alignment=TA_RIGHT,
        spaceBefore=10,
        spaceAfter=6,
    )
    rtl_title = ParagraphStyle(
        'RTLTitle',
        parent=styles['Heading1'],
        fontName=base_font,
        fontSize=16,
        leading=20,
        alignment=TA_RIGHT,
        spaceAfter=12,
    )

    doc = SimpleDocTemplate(PDF_PATH, pagesize=A4, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    flow = [Paragraph("קמילהום – תרשים זרימה ואפיון UI מפורט", rtl_title),
            Paragraph("גרסה: 0.1 • תאריך: 22/08/2025", rtl_style),
            Spacer(1, 12)]

    for title, content in SPEC_SECTIONS:
        flow.append(Paragraph(title, rtl_heading))
        for para in content.split("\n"):
            if para.strip():
                flow.append(Paragraph(para, rtl_style))
        flow.append(Spacer(1, 8))

    doc.build(flow)
    print(f"✅ נוצר PDF: {PDF_PATH}")
    if base_font == "Helvetica":
        print("ℹ️ הערה: כדי לוודא עברית תקינה ב‑PDF, מומלץ להתקין גופן עם עברית (למשל DejaVu Sans) ולעדכן את הנתיב ב‑preferred_fonts.")

# ==== 3) יצירת ויירפריימים (matplotlib) ====
def build_wireframes():
    try:
        import matplotlib.pyplot as plt
        from matplotlib.patches import Rectangle
    except ImportError:
        print("❌ חסרה חבילה: matplotlib. התקן עם: pip install matplotlib")
        sys.exit(1)

    def box(ax, xy, wh, label=None, fontsize=9):
        r = Rectangle(xy, wh[0], wh[1], fill=False)
        ax.add_patch(r)
        if label:
            ax.text(xy[0]+wh[0]/2, xy[1]+wh[1]/2, label, ha="center", va="center", fontsize=fontsize)

    def make_fig(title, size=(7,5)):
        fig, ax = plt.subplots(figsize=size)
        ax.set_title(title, fontsize=14, fontweight="bold")
        ax.axis("off")
        return fig, ax

    # 1) Dashboard
    fig, ax = make_fig("Wireframe – Dashboard")
    box(ax, (0.05,0.7), (0.9,0.25), "אירועים קרובים")
    box(ax, (0.05,0.4), (0.4,0.25), "משפך לידים")
    box(ax, (0.55,0.4), (0.4,0.25), "גבייה צפויה")
    box(ax, (0.05,0.1), (0.9,0.25), "התראות חוזים/גבייה")
    fig.savefig(os.path.join(OUT_DIR, "wireframe_dashboard.png"), bbox_inches="tight"); plt.close(fig)

    # 2) Event Form
    fig, ax = make_fig("Wireframe – Event Form", (6,6))
    y = 0.9
    fields = ["לקוח","תאריך ושעה","סוג אירוע","חבילה","כמות ילדים","כמות מבוגרים","מחיר חבילה","מחיר מנה","סה\"כ","הערות"]
    for f in fields:
        box(ax,(0.1,y-0.05),(0.8,0.06),f); y-=0.08
    box(ax,(0.1,0.05),(0.35,0.06),"שמור"); box(ax,(0.55,0.05),(0.35,0.06),"שלח חוזה")
    fig.savefig(os.path.join(OUT_DIR, "wireframe_event_form.png"), bbox_inches="tight"); plt.close(fig)

    # 3) Customers List
    fig, ax = make_fig("Wireframe – Customers List")
    box(ax,(0.05,0.85),(0.9,0.1),"חיפוש + פילטרים")
    box(ax,(0.05,0.1),(0.9,0.7),"טבלת לקוחות (שם | טלפון | אימייל | אירועים | פעולות)")
    box(ax,(0.05,0.03),(0.2,0.05),"לקוח חדש")
    fig.savefig(os.path.join(OUT_DIR, "wireframe_customers_list.png"), bbox_inches="tight"); plt.close(fig)

    # 4) Customer Form
    fig, ax = make_fig("Wireframe – Customer Form", (6,6))
    y = 0.85
    for f in ["שם מלא","טלפון","אימייל","מקור ליד","כתובת","הערות"]:
        box(ax,(0.1,y),(0.8,0.07),f); y-=0.09
    box(ax,(0.1,0.1),(0.3,0.07),"שמור"); box(ax,(0.5,0.1),(0.3,0.07),"פתח אירוע")
    fig.savefig(os.path.join(OUT_DIR, "wireframe_customer_form.png"), bbox_inches="tight"); plt.close(fig)

    # 5) Packages Catalog
    fig, ax = make_fig("Wireframe – Packages Catalog")
    for i in range(3):
        box(ax,(0.05+0.31*i,0.55),(0.28,0.35),f"חבילה {i+1}\nמחיר בסיס\nמה כלול\n[עריכה] [ארכוב]")
    for i in range(3):
        box(ax,(0.05+0.31*i,0.1),(0.28,0.35),f"חבילה {i+4}\nמחיר בסיס\nמה כלול\n[עריכה] [ארכוב]")
    box(ax,(0.05,0.92),(0.2,0.06),"הוספת חבילה")
    fig.savefig(os.path.join(OUT_DIR, "wireframe_packages_catalog.png"), bbox_inches="tight"); plt.close(fig)

    # 6) Contract Templates
    fig, ax = make_fig("Wireframe – Contract Templates")
    box(ax,(0.05,0.8),(0.9,0.1),"תבנית נבחרת: [ברירת מחדל ▼]  |  משתני Merge")
    box(ax,(0.05,0.1),(0.6,0.65),"עורך תוכן WYSIWYG")
    box(ax,(0.68,0.55),(0.27,0.2),"משתנים\n{{customer.name}}\n{{event.date}}")
    box(ax,(0.68,0.3),(0.27,0.2),"תצוגה מקדימה")
    box(ax,(0.68,0.1),(0.27,0.15),"שמירה / פרסום")
    fig.savefig(os.path.join(OUT_DIR, "wireframe_contract_templates.png"), bbox_inches="tight"); plt.close(fig)

    # 7) Contract Preview & Sign
    fig, ax = make_fig("Wireframe – Contract Preview & Sign")
    box(ax,(0.05,0.75),(0.9,0.2),"פרטי אירוע (לקוח, תאריך, חבילה, סה\"כ)")
    box(ax,(0.05,0.15),(0.9,0.55),"תצוגה מקדימה של החוזה (PDF)")
    box(ax,(0.05,0.05),(0.25,0.07),"שליחה לחתימה")
    box(ax,(0.35,0.05),(0.25,0.07),"הורדת PDF")
    box(ax,(0.65,0.05),(0.25,0.07),"סטטוס חתימה")
    fig.savefig(os.path.join(OUT_DIR, "wireframe_contract_preview.png"), bbox_inches="tight"); plt.close(fig)

    # 8) Reports
    fig, ax = make_fig("Wireframe – Reports")
    box(ax,(0.05,0.85),(0.9,0.1),"סינון: טווח תאריכים, חבילה, סוג אירוע, נציג")
    box(ax,(0.05,0.55),(0.9,0.25),"גרף הכנסות חודשי")
    box(ax,(0.05,0.1),(0.42,0.4),"טבלת הכנסות")
    box(ax,(0.53,0.1),(0.42,0.4),"דו\"ח המרות ומשפך")
    fig.savefig(os.path.join(OUT_DIR, "wireframe_reports.png"), bbox_inches="tight"); plt.close(fig)

    # 9) Settings & Roles
    fig, ax = make_fig("Wireframe – Settings & Roles")
    box(ax,(0.05,0.8),(0.9,0.15),"פרטי עסק + לוגו + מע\"מ")
    box(ax,(0.05,0.5),(0.42,0.25),"תפקידים והרשאות\nמנהל | מכירות")
    box(ax,(0.53,0.5),(0.42,0.25),"תבניות מייל/SMS")
    box(ax,(0.05,0.1),(0.42,0.35),"אינטגרציות חתימה/סליקה")
    box(ax,(0.53,0.1),(0.42,0.35),"Audit & לוגים")
    fig.savefig(os.path.join(OUT_DIR, "wireframe_settings_roles.png"), bbox_inches="tight"); plt.close(fig)

    # 10) Events List
    fig, ax = make_fig("Wireframe – Events List")
    box(ax,(0.05,0.85),(0.9,0.1),"חיפוש + פילטרים (תאריך, סטטוס, סוג אירוע, חבילה)")
    box(ax,(0.05,0.1),(0.9,0.7),"טבלת אירועים (תאריך | לקוח | סוג | סטטוס | סה\"כ | פעולות)")
    box(ax,(0.05,0.03),(0.2,0.05),"אירוע חדש")
    fig.savefig(os.path.join(OUT_DIR, "wireframe_events_list.png"), bbox_inches="tight"); plt.close(fig)

    # 11) Event Details
    fig, ax = make_fig("Wireframe – Event Details")
    box(ax,(0.05,0.7),(0.9,0.25),"כרטיס אירוע: פרטים, חבילה, תמחור, הערות")
    box(ax,(0.05,0.4),(0.45,0.25),"פעולות: הצעה, חוזה, תשלום, שינוי סטטוס")
    box(ax,(0.55,0.4),(0.4,0.25),"לוג שינויים / גרסאות")
    box(ax,(0.05,0.1),(0.9,0.25),"מסמכים מצורפים (חוזה, הצעות, חשבוניות)")
    fig.savefig(os.path.join(OUT_DIR, "wireframe_event_details.png"), bbox_inches="tight"); plt.close(fig)

    # 12) Payments
    fig, ax = make_fig("Wireframe – Payments")
    box(ax,(0.05,0.8),(0.9,0.15),"קישור סליקה / תשלום ידני / מקדמה")
    box(ax,(0.05,0.5),(0.42,0.25),"תשלומים שבוצעו")
    box(ax,(0.53,0.5),(0.42,0.25),"יתרות פתוחות")
    box(ax,(0.05,0.1),(0.9,0.35),"היסטוריית חשבוניות/קבלות")
    fig.savefig(os.path.join(OUT_DIR, "wireframe_payments.png"), bbox_inches="tight"); plt.close(fig)

    # 13) Calendar View
    fig, ax = make_fig("Wireframe – Calendar View")
    box(ax,(0.05,0.85),(0.9,0.1),"סינון: תאריך/חודש, סוג אירוע, איש מכירות")
    box(ax,(0.05,0.1),(0.9,0.7),"תצוגת חודש: אירועים לפי ימים")
    fig.savefig(os.path.join(OUT_DIR, "wireframe_calendar.png"), bbox_inches="tight"); plt.close(fig)

    print("✅ נוצרו ויירפריימים (PNG) בתיקייה:", OUT_DIR)

# ==== 4) יצירת ZIP ====
def build_zip():
    with zipfile.ZipFile(ZIP_PATH, 'w', zipfile.ZIP_DEFLATED) as zf:
        # מסמכים
        for path in [DOCX_PATH, PDF_PATH]:
            if os.path.exists(path):
                zf.write(path, arcname=os.path.basename(path))
        # תמונות
        for fname, _ in WIREFRAMES:
            fpath = os.path.join(OUT_DIR, fname)
            if os.path.exists(fpath):
                zf.write(fpath, arcname=fname)
    print(f"✅ נוצר ZIP: {ZIP_PATH}")

def main():
    build_docx()
    build_pdf()
    build_wireframes()
    build_zip()
    print("\n🎉 מוכן! כל הקבצים נמצאים בתיקייה:", OUT_DIR)

if __name__ == "__main__":
    main()


